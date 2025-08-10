import os
import json
from typing import Dict, List
from datetime import datetime, timedelta
from jinja2 import Template
import logging
from .models import TranscriptSession, SpeakerSegment

logger = logging.getLogger(__name__)


class TranscriptExporter:
    """Service for exporting transcripts in various formats"""
    
    def __init__(self, export_dir: str = "exports"):
        self.export_dir = export_dir
        os.makedirs(export_dir, exist_ok=True)
    
    def format_timestamp(self, seconds: float) -> str:
        """Format seconds to HH:MM:SS.mmm"""
        td = timedelta(seconds=seconds)
        hours, remainder = divmod(td.seconds, 3600)
        minutes, seconds = divmod(remainder, 60)
        milliseconds = int(td.microseconds / 1000)
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}.{milliseconds:03d}"
    
    def export_as_txt(self, session_data: Dict, speaker_names: Dict[str, str] = None) -> str:
        """Export transcript as plain text"""
        try:
            segments = session_data.get("segments", [])
            session_name = session_data.get("session_name", "Transcript")
            created_at = session_data.get("created_at", datetime.utcnow())
            
            # Template for text export
            txt_template = Template("""
{{ session_name }}
Generated on: {{ created_at.strftime('%Y-%m-%d %H:%M:%S') }}
Total Segments: {{ segments|length }}

{% for segment in segments -%}
[{{ format_timestamp(segment.start_time) }} - {{ format_timestamp(segment.end_time) }}] {{ speaker_name(segment.speaker_id) }}: {{ segment.text }}
{% endfor %}
""".strip())
            
            def speaker_name(speaker_id):
                if speaker_names and speaker_id in speaker_names:
                    return speaker_names[speaker_id]
                return speaker_id
            
            content = txt_template.render(
                session_name=session_name,
                created_at=created_at,
                segments=segments,
                format_timestamp=self.format_timestamp,
                speaker_name=speaker_name
            )
            
            # Save to file
            filename = f"{session_data.get('_id', 'transcript')}_{int(datetime.utcnow().timestamp())}.txt"
            filepath = os.path.join(self.export_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Exported transcript as TXT: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error exporting as TXT: {e}")
            raise
    
    def export_as_srt(self, session_data: Dict, speaker_names: Dict[str, str] = None) -> str:
        """Export transcript as SRT subtitle file"""
        try:
            segments = session_data.get("segments", [])
            
            def speaker_name(speaker_id):
                if speaker_names and speaker_id in speaker_names:
                    return speaker_names[speaker_id]
                return speaker_id
            
            # SRT format
            srt_content = []
            
            for i, segment in enumerate(segments, 1):
                start_time = self.format_timestamp(segment.get("start_time", 0))
                end_time = self.format_timestamp(segment.get("end_time", 0))
                speaker = speaker_name(segment.get("speaker_id", "Unknown"))
                text = segment.get("text", "").strip()
                
                if not text:
                    continue
                
                srt_content.append(f"{i}")
                srt_content.append(f"{start_time} --> {end_time}")
                srt_content.append(f"{speaker}: {text}")
                srt_content.append("")  # Empty line between entries
            
            content = "\n".join(srt_content)
            
            # Save to file
            filename = f"{session_data.get('_id', 'transcript')}_{int(datetime.utcnow().timestamp())}.srt"
            filepath = os.path.join(self.export_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Exported transcript as SRT: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error exporting as SRT: {e}")
            raise
    
    def export_as_json(self, session_data: Dict, speaker_names: Dict[str, str] = None) -> str:
        """Export transcript as JSON"""
        try:
            # Prepare JSON structure
            export_data = {
                "session_info": {
                    "id": str(session_data.get("_id", "")),
                    "name": session_data.get("session_name", ""),
                    "created_at": session_data.get("created_at").isoformat() if session_data.get("created_at") else None,
                    "updated_at": session_data.get("updated_at").isoformat() if session_data.get("updated_at") else None,
                    "language": session_data.get("language"),
                    "status": session_data.get("status", ""),
                    "total_duration": session_data.get("total_duration", 0.0)
                },
                "speakers": speaker_names or {},
                "segments": [],
                "export_info": {
                    "exported_at": datetime.utcnow().isoformat(),
                    "format_version": "1.0",
                    "total_segments": len(session_data.get("segments", []))
                }
            }
            
            # Process segments
            for segment in session_data.get("segments", []):
                segment_data = {
                    "speaker_id": segment.get("speaker_id", ""),
                    "speaker_name": speaker_names.get(segment.get("speaker_id", "")) if speaker_names else None,
                    "start_time": segment.get("start_time", 0.0),
                    "end_time": segment.get("end_time", 0.0),
                    "duration": segment.get("end_time", 0.0) - segment.get("start_time", 0.0),
                    "text": segment.get("text", "").strip(),
                    "confidence": segment.get("confidence", 0.0),
                    "timestamp": segment.get("timestamp").isoformat() if segment.get("timestamp") else None
                }
                export_data["segments"].append(segment_data)
            
            # Save to file
            filename = f"{session_data.get('_id', 'transcript')}_{int(datetime.utcnow().timestamp())}.json"
            filepath = os.path.join(self.export_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Exported transcript as JSON: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error exporting as JSON: {e}")
            raise
    
    def export_as_csv(self, session_data: Dict, speaker_names: Dict[str, str] = None) -> str:
        """Export transcript as CSV"""
        try:
            import csv
            
            segments = session_data.get("segments", [])
            
            def speaker_name(speaker_id):
                if speaker_names and speaker_id in speaker_names:
                    return speaker_names[speaker_id]
                return speaker_id
            
            # Save to file
            filename = f"{session_data.get('_id', 'transcript')}_{int(datetime.utcnow().timestamp())}.csv"
            filepath = os.path.join(self.export_dir, filename)
            
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                
                # Header
                writer.writerow([
                    'Segment',
                    'Speaker ID',
                    'Speaker Name',
                    'Start Time',
                    'End Time',
                    'Duration',
                    'Text',
                    'Confidence'
                ])
                
                # Data rows
                for i, segment in enumerate(segments, 1):
                    start_time = segment.get("start_time", 0.0)
                    end_time = segment.get("end_time", 0.0)
                    duration = end_time - start_time
                    speaker_id = segment.get("speaker_id", "")
                    
                    writer.writerow([
                        i,
                        speaker_id,
                        speaker_name(speaker_id),
                        self.format_timestamp(start_time),
                        self.format_timestamp(end_time),
                        f"{duration:.2f}s",
                        segment.get("text", "").strip(),
                        f"{segment.get('confidence', 0.0):.2f}"
                    ])
            
            logger.info(f"Exported transcript as CSV: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error exporting as CSV: {e}")
            raise
    
    def export_as_docx(self, session_data: Dict, speaker_names: Dict[str, str] = None) -> str:
        """Export transcript as Word document (requires python-docx)"""
        try:
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                raise ImportError("python-docx is required for DOCX export")
            
            segments = session_data.get("segments", [])
            session_name = session_data.get("session_name", "Transcript")
            created_at = session_data.get("created_at", datetime.utcnow())
            
            def speaker_name(speaker_id):
                if speaker_names and speaker_id in speaker_names:
                    return speaker_names[speaker_id]
                return speaker_id
            
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading(session_name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Generated on: {created_at.strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Total Segments: {len(segments)}")
            doc.add_paragraph("")  # Empty line
            
            # Segments
            for segment in segments:
                start_time = self.format_timestamp(segment.get("start_time", 0))
                end_time = self.format_timestamp(segment.get("end_time", 0))
                speaker = speaker_name(segment.get("speaker_id", "Unknown"))
                text = segment.get("text", "").strip()
                confidence = segment.get("confidence", 0.0)
                
                if not text:
                    continue
                
                # Time and speaker info
                p = doc.add_paragraph()
                p.add_run(f"[{start_time} - {end_time}] ").bold = True
                p.add_run(f"{speaker}: ").bold = True
                p.add_run(text)
                
                if confidence < 0.7:  # Low confidence indicator
                    p.add_run(f" (Low confidence: {confidence:.2f})").italic = True
            
            # Save to file
            filename = f"{session_data.get('_id', 'transcript')}_{int(datetime.utcnow().timestamp())}.docx"
            filepath = os.path.join(self.export_dir, filename)
            
            doc.save(filepath)
            
            logger.info(f"Exported transcript as DOCX: {filepath}")
            return filepath
            
        except Exception as e:
            logger.error(f"Error exporting as DOCX: {e}")
            raise
    
    def export_transcript(self, session_data: Dict, format_type: str, speaker_names: Dict[str, str] = None) -> str:
        """Export transcript in specified format"""
        format_type = format_type.lower()
        
        export_methods = {
            'txt': self.export_as_txt,
            'text': self.export_as_txt,
            'srt': self.export_as_srt,
            'json': self.export_as_json,
            'csv': self.export_as_csv,
            'docx': self.export_as_docx,
            'word': self.export_as_docx
        }
        
        if format_type not in export_methods:
            raise ValueError(f"Unsupported export format: {format_type}")
        
        return export_methods[format_type](session_data, speaker_names)


# Global exporter instance
exporter = TranscriptExporter()
